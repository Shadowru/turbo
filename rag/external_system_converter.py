import docx
import datetime
import json
import os

# === НАСТРОЙКИ ===
INPUT_FILE = "D:/Work/Python/turbo/data/Продукты МЭШ-v1035-20251210_210751.docx"  # Имя вашего файла Word с таблицей
OUTPUT_FILE = 'system_registry.md' # Имя выходного файла

def clean_text(text):
    """Очистка текста от лишних пробелов."""
    if text:
        return text.strip()
    return ""

def extract_table_data(docx_path):
    """Читает первую таблицу из docx и превращает её в список словарей."""
    if not os.path.exists(docx_path):
        print(f"Ошибка: Файл {docx_path} не найден.")
        return []

    doc = docx.Document(docx_path)
    
    if not doc.tables:
        print("Ошибка: В документе не найдено таблиц.")
        return []

    # Берем первую таблицу
    table = doc.tables[1]
    
    # Получаем заголовки из первой строки
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    
    data = []
    
    # Итерируемся по строкам, начиная со второй
    for row in table.rows[1:]:
        row_data = {}
        # Пропускаем пустые строки
        if not any(cell.text.strip() for cell in row.cells):
            continue
            
        for i, cell in enumerate(row.cells):
            if i < len(headers):
                header_name = headers[i]
                row_data[header_name] = clean_text(cell.text)
        
        data.append(row_data)
        
    return data

def generate_markdown(data):
    """Генерирует текст в нужном формате на основе данных."""
    
    # Глобальный заголовок файла (Frontmatter)
    current_date = datetime.date.today().strftime("%Y-%m-%d")
    output = []
    output.append("---")
    output.append("document_type: system_registry")
    output.append("version: 1.2")
    output.append(f"updated_at: {current_date}")
    output.append("---\n")

    for item in data:
        # Извлечение полей с защитой от отсутствия ключа
        product_code = item.get('Product code', 'NO_CODE')
        product_name = item.get('Продукт', 'Unnamed System')
        description = item.get('Назначение', 'Описание отсутствует.')
        owner = item.get('Владелец', '-')
        status = item.get('Статус', '-')
        
        # Формирование алиасов (добавляем Код КЕ и имя продукта)
        aliases = []
        if item.get('Код КЕ'):
            aliases.append(item.get('Код КЕ'))
        aliases_json = json.dumps(aliases, ensure_ascii=False)

        # Ссылки и ресурсы (Jira, Wiki, Repo)
        jira = item.get('Jira', '-')
        wiki = item.get('Wiki', '-')
        repo = item.get('Repo', '-')
        
        # --- ГЕНЕРАЦИЯ КАРТОЧКИ СИСТЕМЫ ---
        
        output.append(f"## System Card: {product_code}")
        output.append(f"- **system_id:** {product_code}")
        output.append(f"- **system_name:** {product_name}")
        output.append(f"- **aliases:** {aliases_json}")
        output.append(f"- **status:** {status}") # Добавлено из таблицы
        output.append("")
        
        output.append("### Назначение")
        output.append(description)
        output.append("")
        
        # Поскольку в исходной таблице нет колонок "Основные функции" и "Сервисы",
        # мы создадим раздел "Ресурсы и Детали" на основе имеющихся данных (Jira, Wiki, Repo и т.д.)
        
        output.append("### Ресурсы и ссылки")
        output.append(f"| Ресурс | Ссылка/Значение |")
        output.append(f"|---|---|")
        output.append(f"| Jira | {jira} |")
        output.append(f"| Wiki | {wiki} |")
        output.append(f"| Repo | {repo} |")
        output.append(f"| Проект мониторинга | {item.get('Проект мониторинга', '-')} |")
        output.append("")

        output.append("### Жизненный цикл и SLA")
        output.append(f"- **Приоритет:** {item.get('Приоритет продукта', '-')}")
        output.append(f"- **ЗНИ на приемку:** {item.get('ЗНИ на приемку продукты в эксплуатацию', '-')}")
        output.append(f"- **ЗНИ на вывод:** {item.get('ЗНИ на вывод из эксплуатации', '-')}")
        output.append(f"- **ФГИС/Регион:** {item.get('ФГИС/Регион', '-')}")
        output.append("")

        output.append("### Контакты и владельцы")
        output.append(f"Владелец: {owner}")
        output.append(f"Продуктовое направление: {item.get('Продуктовое направление', '-')}")
        
        output.append("\n---\n")

    return "\n".join(output)

def main():
    print(f"Чтение файла {INPUT_FILE}...")
    data = extract_table_data(INPUT_FILE)
    
    if not data:
        print("Нет данных для обработки.")
        return

    print(f"Найдено записей: {len(data)}")
    print("Генерация Markdown...")
    
    markdown_content = generate_markdown(data)
    
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
        
    print(f"Готово! Результат сохранен в {OUTPUT_FILE}")

if __name__ == "__main__":
    main()